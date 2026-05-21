"""
Document loading utilities for PDF, DOCX, and TXT files.
"""

import io
import os
import logging
from pathlib import Path
from typing import List, Optional, Dict, Any

logger = logging.getLogger(__name__)


def load_pdf(file_obj) -> List[Dict[str, Any]]:
    """Load and extract text from a PDF file using PyPDF2 with fallback to pypdf."""
    documents = []
    file_name = getattr(file_obj, "name", "unknown.pdf")
    file_bytes = file_obj.read() if hasattr(file_obj, "read") else file_obj

    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                documents.append({
                    "page_content": text.strip(),
                    "metadata": {
                        "source": file_name,
                        "page": page_num + 1,
                        "file_type": "pdf",
                        "total_pages": len(reader.pages),
                    },
                })
    except Exception as e1:
        logger.warning(f"PyPDF2 failed ({e1}), trying pypdf...")
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    documents.append({
                        "page_content": text.strip(),
                        "metadata": {
                            "source": file_name,
                            "page": page_num + 1,
                            "file_type": "pdf",
                            "total_pages": len(reader.pages),
                        },
                    })
        except Exception as e2:
            raise RuntimeError(f"Failed to load PDF: {e1} | {e2}") from e2

    return documents


def load_docx(file_obj) -> List[Dict[str, Any]]:
    """Load and extract text from a DOCX file."""
    from docx import Document
    file_name = getattr(file_obj, "name", "unknown.docx")
    file_bytes = file_obj.read() if hasattr(file_obj, "read") else file_obj

    doc = Document(io.BytesIO(file_bytes))
    full_text = "\n".join(
        para.text for para in doc.paragraphs if para.text.strip()
    )

    if not full_text.strip():
        return []

    return [{
        "page_content": full_text.strip(),
        "metadata": {
            "source": file_name,
            "file_type": "docx",
            "paragraph_count": len(doc.paragraphs),
        },
    }]


def load_txt(file_obj) -> List[Dict[str, Any]]:
    """Load and extract text from a TXT file."""
    file_name = getattr(file_obj, "name", "unknown.txt")
    raw = file_obj.read() if hasattr(file_obj, "read") else file_obj
    if isinstance(raw, bytes):
        text = raw.decode("utf-8", errors="replace")
    else:
        text = raw

    if not text.strip():
        return []

    return [{
        "page_content": text.strip(),
        "metadata": {
            "source": file_name,
            "file_type": "txt",
            "char_count": len(text),
        },
    }]


def load_document(file_obj) -> List[Dict[str, Any]]:
    """Dispatch loader based on file extension."""
    file_name = getattr(file_obj, "name", "")
    ext = Path(file_name).suffix.lower()

    loaders = {
        ".pdf": load_pdf,
        ".docx": load_docx,
        ".doc": load_docx,
        ".txt": load_txt,
        ".md": load_txt,
    }

    loader = loaders.get(ext)
    if loader is None:
        raise ValueError(f"Unsupported file type: '{ext}'. Supported: {list(loaders.keys())}")

    docs = loader(file_obj)
    logger.info(f"Loaded {len(docs)} document section(s) from '{file_name}'.")
    return docs


def load_multiple_documents(files: list) -> List[Dict[str, Any]]:
    """Load multiple files and aggregate the results."""
    all_docs = []
    errors = []

    for f in files:
        try:
            docs = load_document(f)
            all_docs.extend(docs)
        except Exception as e:
            errors.append(f"{getattr(f, 'name', 'file')}: {e}")
            logger.error(f"Error loading file: {e}")

    if errors:
        logger.warning(f"Errors during loading: {errors}")

    return all_docs, errors
